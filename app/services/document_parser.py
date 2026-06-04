"""文档解析服务：Word/PDF → Markdown。"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any


def parse_document(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """解析文档并返回 Markdown 内容。

    Args:
        file_bytes: 文件二进制内容
        filename: 原始文件名（用于判断格式和推断标题）

    Returns:
        {"title": str, "content": str, "summary": str}
    """
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        return _parse_docx(file_bytes, filename)
    elif ext == ".pdf":
        return _parse_pdf(file_bytes, filename)
    else:
        raise ValueError(f"不支持的文件格式: {ext}（仅支持 .docx / .pdf）")


def _parse_docx(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """解析 Word 文档为 Markdown。"""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    title = Path(filename).stem  # 默认用文件名作标题

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = (para.style.name or "").lower()

        # 从标题样式推断 Markdown 标题级别
        if "heading 1" in style_name or "标题 1" in style_name:
            lines.append(f"# {text}")
            if i < 3:
                title = text  # 用第一个 H1 作为文档标题
        elif "heading 2" in style_name or "标题 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name or "标题 3" in style_name:
            lines.append(f"### {text}")
        elif "heading 4" in style_name or "标题 4" in style_name:
            lines.append(f"#### {text}")
        elif "list" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    content = "\n\n".join(lines)
    # 清理多余空行
    content = re.sub(r"\n{3,}", "\n\n", content)

    # 提取表格
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        if table_md:
            content += f"\n\n{table_md}"

    summary = _extract_summary(content)
    return {"title": title, "content": content, "summary": summary}


def _parse_pdf(file_bytes: bytes, filename: str) -> dict[str, Any]:
    """解析 PDF 文档为 Markdown。"""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    title = Path(filename).stem
    pages: list[str] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue

        # 第一页尝试提取标题（取第一行非空文本）
        if i == 0:
            first_line = text.split("\n")[0].strip()
            if first_line and len(first_line) < 100:
                title = first_line

        pages.append(f"<!-- 第 {i + 1} 页 -->\n\n{text}")

    content = "\n\n---\n\n".join(pages)
    summary = _extract_summary(content)
    return {"title": title, "content": content, "summary": summary}


def _table_to_markdown(table) -> str:
    """将 docx 表格转为 Markdown 表格。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # 构建 Markdown 表格
    lines = []
    # 表头
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # 数据行
    for row in rows[1:]:
        # 补齐列数
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")

    return "\n".join(lines)


def _extract_summary(content: str, max_len: int = 200) -> str:
    """从内容中提取摘要（前 N 个字符的纯文本）。"""
    # 去掉 Markdown 标记
    text = re.sub(r"[#*\-|`>]", "", content)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > max_len:
        return text[:max_len] + "..."
    return text


def auto_categorize(title: str, content: str) -> str:
    """根据标题和内容自动判断文档分类。

    分类规则：
    - competitive: 含竞品/对手/友商关键词（AWS/阿里/华为/Cloudflare/竞争/调研/对比）
    - sop: 含 SOP/流程/步骤/操作指南
    - manual: 含产品/架构/背景/规划/开区
    - technical: 含 API/接口/代码/配置/部署
    - 默认: manual
    """
    text = (title + " " + content[:2000]).lower()

    # 竞争分析关键词
    competitive_keywords = [
        "aws", "阿里", "华为", "cloudflare", "竞争", "竞分", "友商",
        "对比", "调研", "ens", "local zone", "wavelength", "对手",
        "市场分析", "行业分析", "营收", "定价对比",
    ]
    if any(kw in text for kw in competitive_keywords):
        return "competitive"

    # SOP 流程
    sop_keywords = ["sop", "流程", "步骤", "操作指南", "checklist", "工单流"]
    if any(kw in text for kw in sop_keywords):
        return "sop"

    # 技术文档
    tech_keywords = ["api", "接口文档", "sdk", "部署", "配置", "代码", "开发指南", "架构设计"]
    if any(kw in text for kw in tech_keywords):
        return "technical"

    # 默认归入运营手册
    return "manual"
